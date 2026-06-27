"""
Ingest documents from config.DOCS_DIR into a Qdrant hybrid vector store.

Pipeline:
  [0/3]  Convert legacy office files (.xls/.doc/.ppt) into modern formats
  [0b/3] Convert any WMF/EMF vector images inside Office zips to PNG
  [1/3]  Parse all files with Docling in parallel (OCR disabled for speed)
  [1b/3] Re-run scanned PDFs that produced 0 chunks, this time with OCR
  [2/3]  Load the BGE-M3 embedding model
  [3/3]  Embed all chunks and index them into Qdrant

Usage:
    python create_database.py
"""

import io
import shutil
import subprocess
import zipfile
from concurrent.futures import as_completed
from pathlib import Path

import transformers
transformers.logging.set_verbosity_error()   # silence HuggingFace progress spam

import tiktoken
from docling.datamodel.base_models import DocItemLabel, InputFormat
from docling.datamodel.document import DoclingDocument, PictureItem
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling_core.transforms.chunker import HybridChunker
from docling_core.transforms.chunker.tokenizer.openai import OpenAITokenizer
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import FastEmbedSparse, QdrantVectorStore, RetrievalMode
from langchain_text_splitters import RecursiveCharacterTextSplitter

import config

# Optional XLS support — only needed for legacy .xls files
try:
    import xlrd
    import openpyxl
    _HAVE_XLS = True
except ImportError:
    _HAVE_XLS = False


# ── Constants ─────────────────────────────────────────────────────────────────

DOCS_DIR        = Path(config.DOCS_DIR)          # source documents folder
CONVERTED_CACHE = Path("./data/converted_cache") # all files land here before parsing

# File extensions that Docling can parse directly
SUPPORTED_MODERN = {".pdf", ".docx", ".pptx", ".xlsx", ".txt", ".md"}

# Docling labels that indicate a visual element (figure or chart)
_VISUAL_LABELS = {DocItemLabel.PICTURE, DocItemLabel.CHART}

# Vector image formats embedded inside Office zip files
_WMF_SUFFIXES = frozenset({".wmf", ".emf"})

# Office formats that are ZIP archives (and may contain WMF/EMF images inside)
_OFFICE_ZIPS = frozenset({".docx", ".pptx", ".pptm", ".ppsx", ".docm", ".xlsm"})

# Number of parallel Docling workers. Each worker uses ~2 GB RAM.
# Lower this if you run out of memory.
NUM_WORKERS = 6

# Temp directory for LibreOffice EMF conversion output
_VECTOR_TMP = Path("./data/.vector_tmp").resolve()


# ── WMF/EMF pre-conversion ────────────────────────────────────────────────────
# Office files (docx, pptx, etc.) are ZIP archives that can contain WMF or EMF
# vector images. Docling can't render these, so we convert them to PNG first
# by rewriting the ZIP in-place before handing the file to Docling.

def _rewrite_zip_with_pngs(
    zip_path: Path,
    rename_map: dict[str, str],
    png_data: dict[str, bytes],
) -> None:
    """
    Rewrite an Office ZIP file, replacing WMF/EMF entries with pre-rendered PNGs.

    rename_map: old zip entry name → new zip entry name (with .png extension)
    png_data:   new zip entry name → PNG bytes to write
    Also patches the .xml and .rels files inside the ZIP so they reference
    the new PNG filenames and content types instead of the old vector ones.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(zip_path, "r") as zin:
        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                if info.filename in rename_map:
                    continue  # this entry is replaced by a PNG below
                raw = zin.read(info.filename)
                # Patch XML/RELS to reference new filenames and content types
                if info.filename.endswith((".xml", ".rels")):
                    text = raw.decode("utf-8", errors="replace")
                    for old, new in rename_map.items():
                        text = text.replace(Path(old).name, Path(new).name)
                    text = (text
                            .replace('ContentType="image/x-wmf"', 'ContentType="image/png"')
                            .replace('ContentType="image/x-emf"', 'ContentType="image/png"'))
                    raw = text.encode("utf-8")
                zout.writestr(info, raw)
            # Write the new PNG entries
            for new_name, data in png_data.items():
                zout.writestr(new_name, data)
    zip_path.write_bytes(buf.getvalue())


def _batch_convert_vectors_in_cache() -> int:
    """
    Find all WMF/EMF images inside cached Office ZIPs and convert them to PNG.

    Two converters are tried in order:
    - WMF files → ImageMagick `convert` (fast, per-file)
    - EMF files and any WMF that ImageMagick failed on → LibreOffice (one batch call)

    After conversion, each ZIP is rewritten in-place with PNGs replacing the
    original vector images. Returns the total number of images converted.
    """
    _VECTOR_TMP.mkdir(parents=True, exist_ok=True)

    # Build inventory: zip_path → {entry_name → (suffix, raw_bytes)}
    inventory: dict[Path, dict[str, tuple[str, bytes]]] = {}
    for zip_path in CONVERTED_CACHE.iterdir():
        if zip_path.suffix.lower() not in _OFFICE_ZIPS:
            continue
        try:
            with zipfile.ZipFile(zip_path, "r") as zin:
                vectors = [n for n in zin.namelist() if Path(n).suffix.lower() in _WMF_SUFFIXES]
                if vectors:
                    inventory[zip_path] = {
                        n: (Path(n).suffix.lower().lstrip("."), zin.read(n))
                        for n in vectors
                    }
        except Exception:
            pass

    if not inventory:
        return 0

    import tempfile, os

    # ── Step 1: try ImageMagick for WMF files ─────────────────────────────────
    wmf_pngs:  dict[bytes, bytes] = {}   # raw bytes → PNG bytes (deduplicated by content)
    wmf_failed: set[bytes]        = set() # raws ImageMagick couldn't handle

    all_wmf_raws = {
        raw
        for entries in inventory.values()
        for suffix, raw in entries.values()
        if suffix == "wmf"
    }
    for raw in all_wmf_raws:
        fd, tmp_in = tempfile.mkstemp(suffix=".wmf")
        try:
            os.write(fd, raw)
            os.close(fd)
            fd = -1
            tmp_out = tmp_in[:-4] + ".png"
            res = subprocess.run(
                ["convert", "-density", "150", "-background", "white", tmp_in, tmp_out],
                capture_output=True, timeout=30,
            )
            if res.returncode == 0 and os.path.exists(tmp_out):
                wmf_pngs[raw] = Path(tmp_out).read_bytes()
                os.unlink(tmp_out)
            else:
                wmf_failed.add(raw)  # may actually be an EMF — retry via LibreOffice
        except Exception:
            wmf_failed.add(raw)
        finally:
            if fd >= 0:
                try:
                    os.close(fd)
                except OSError:
                    pass
            try:
                os.unlink(tmp_in)
            except OSError:
                pass

    # ── Step 2: batch-convert EMF files + failed WMFs via LibreOffice ─────────
    lo_pngs:  dict[bytes, bytes]            = {}
    lo_batch: list[tuple[Path, bytes, str]] = []  # (tmp_file, raw, orig_suffix)
    seen_lo:  set[bytes]                    = set()

    for entries in inventory.values():
        for suffix, raw in entries.values():
            if raw in seen_lo:
                continue
            if suffix == "emf" or (suffix == "wmf" and raw in wmf_failed):
                seen_lo.add(raw)
                stem  = f"vec_{len(lo_batch):06d}"
                tmp_f = _VECTOR_TMP / f"{stem}.emf"  # LibreOffice needs .emf extension
                tmp_f.write_bytes(raw)
                lo_batch.append((tmp_f, raw, suffix))

    if lo_batch:
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "png",
                 "--outdir", str(_VECTOR_TMP)] + [str(p) for p, _, _ in lo_batch],
                capture_output=True, timeout=300,
            )
            for tmp_f, raw, _ in lo_batch:
                out_png = tmp_f.with_suffix(".png")
                if out_png.exists() and out_png.stat().st_size > 100:
                    lo_pngs[raw] = out_png.read_bytes()
        except Exception:
            pass
        finally:
            for tmp_f, _, _ in lo_batch:
                try:
                    tmp_f.unlink()
                except OSError:
                    pass
            for p in _VECTOR_TMP.glob("*.png"):
                try:
                    p.unlink()
                except OSError:
                    pass

    # ── Step 3: rewrite each affected ZIP in-place ────────────────────────────
    all_pngs = {**wmf_pngs, **lo_pngs}
    total    = 0

    for zip_path, entries in inventory.items():
        rename_map: dict[str, str]   = {}
        png_data:   dict[str, bytes] = {}

        for old_name, (suffix, raw) in entries.items():
            png = all_pngs.get(raw)
            if png:
                new_name            = str(Path(old_name).with_suffix(".png"))
                rename_map[old_name] = new_name
                png_data[new_name]   = png

        if rename_map:
            try:
                _rewrite_zip_with_pngs(zip_path, rename_map, png_data)
                total += len(rename_map)
            except Exception as e:
                print(f"      Failed to rewrite {zip_path.name}: {e}", flush=True)

    try:
        _VECTOR_TMP.rmdir()
    except OSError:
        pass

    return total


# ── Legacy format conversion ───────────────────────────────────────────────────
# Old Office formats (.xls, .doc, .ppt) and macro-enabled variants can't be
# parsed by Docling directly. We convert them to modern equivalents first.
# All files (legacy and modern) end up in CONVERTED_CACHE before parsing.

def _cache_path(src: Path, new_ext: str | None = None) -> Path:
    """Return the expected cache path for a source file, optionally with a new extension."""
    ext = new_ext if new_ext else src.suffix
    return CONVERTED_CACHE / f"{src.stem}{ext}"


def _convert_legacy(docs_dir: Path) -> None:
    """
    Scan docs_dir for all files and populate CONVERTED_CACHE:

    - .xls files are converted to .xlsx using xlrd + openpyxl (pure Python, no LibreOffice)
    - .doc, .ppt, .pptm, .ppsx, .docm, .xlsm files are converted via LibreOffice headless
    - Modern formats (.pdf, .docx, .pptx, .xlsx, .txt, .md) are copied as-is
    - .txt files are re-encoded to UTF-8 if needed (some may be latin-1)
    - Office temp/lock files (~$filename) are skipped

    Already-cached files are skipped to avoid redundant work on re-runs.
    """
    CONVERTED_CACHE.mkdir(parents=True, exist_ok=True)

    # Convert .xls → .xlsx using pure Python (faster than LibreOffice for spreadsheets)
    for xls in docs_dir.rglob("*.xls"):
        out = _cache_path(xls, ".xlsx")
        if out.exists():
            continue
        if not _HAVE_XLS:
            print(f"      Skipping {xls.name}: xlrd/openpyxl not installed")
            continue
        try:
            wb_xls  = xlrd.open_workbook(xls)
            wb_xlsx = openpyxl.Workbook()
            for i in range(wb_xls.nsheets):
                sh = wb_xls.sheet_by_index(i)
                ws = wb_xlsx.active if i == 0 else wb_xlsx.create_sheet(title=sh.name)
                ws.title = sh.name
                for row in range(sh.nrows):
                    ws.append(sh.row_values(row))
            wb_xlsx.save(out)
            print(f"      Converted {xls.name} → {out.name}")
        except Exception as e:
            print(f"      Failed {xls.name}: {e}")

    # Convert legacy/macro Office formats via LibreOffice headless
    legacy_map = {
        ".ppt":  ".pptx",
        ".doc":  ".docx",
        ".ppsx": ".pptx",
        ".pptm": ".pptx",
        ".docm": ".docx",
        ".xlsm": ".xlsx",
    }
    legacy_files = [
        p for p in docs_dir.rglob("*")
        if p.is_file() and p.suffix.lower() in legacy_map
    ]
    for src in legacy_files:
        target_ext = legacy_map[src.suffix.lower()]
        out = _cache_path(src, target_ext)
        if out.exists():
            continue
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", target_ext[1:],
                 "--outdir", str(CONVERTED_CACHE.resolve()), str(src.resolve())],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode == 0 and out.exists():
                print(f"      Converted {src.name} → {out.name}")
            else:
                print(f"      Failed {src.name}: {result.stderr.strip() or 'unknown error'}")
        except Exception as e:
            print(f"      Failed {src.name}: {e}")

    # Copy modern formats directly into the cache (no conversion needed)
    for src in docs_dir.rglob("*"):
        if not src.is_file():
            continue
        if src.name.startswith("~$"):          # skip Office temp/lock files
            continue
        if src.suffix.lower() not in SUPPORTED_MODERN:
            continue
        out = _cache_path(src)
        if out.exists():
            continue
        if src.suffix.lower() == ".txt":
            # Re-encode latin-1 text files to UTF-8 so Docling can read them
            try:
                raw = src.read_bytes()
                try:
                    raw.decode("utf-8")
                    out.write_bytes(raw)
                except UnicodeDecodeError:
                    out.write_text(src.read_text(encoding="latin1"), encoding="utf-8")
                    print(f"      Re-encoded {src.name} to UTF-8")
            except Exception as e:
                print(f"      Failed {src.name}: {e}")
        else:
            try:
                shutil.copy2(src, out)
            except Exception as e:
                print(f"      Failed to cache {src.name}: {e}")


# ── Docling converter factory ──────────────────────────────────────────────────

def _make_converter() -> DocumentConverter:
    """
    Create a Docling DocumentConverter with OCR disabled.
    OCR is skipped for speed — scanned PDFs that produce 0 chunks are automatically
    retried with OCR enabled in the main() loop.
    """
    pdf_opts        = PdfPipelineOptions()
    pdf_opts.do_ocr = False
    return DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts)}
    )


# ── Image extraction ───────────────────────────────────────────────────────────

def _extract_images(
    dl_doc: DoclingDocument,
    source_stem: str,
    image_store: Path,
) -> dict[str, str]:
    """
    Save every PictureItem in a parsed document as a PNG file.

    Note: unlike the version in backend.py, this function does NOT filter by
    size or aspect ratio — all images are saved during bulk indexing.
    The backend's live-upload version applies filtering to avoid storing icons
    and decorative elements.

    Returns a dict mapping each item's self_ref ID to its saved file path,
    used by _build_chunks to associate images with their surrounding text.
    """
    image_store.mkdir(parents=True, exist_ok=True)
    ref_to_path: dict[str, str] = {}
    fig_idx = 0

    for item, _level in dl_doc.iterate_items():
        if isinstance(item, PictureItem):
            img = item.get_image(dl_doc)
            if img is not None:
                out_path = image_store / f"{source_stem}__fig{fig_idx:04d}.png"
                # Ensure compatible PIL mode before saving
                if img.mode not in ("RGB", "RGBA", "L", "LA"):
                    img = img.convert("RGBA" if "A" in img.mode else "RGB")
                img.save(out_path, "PNG")
                ref_to_path[item.self_ref] = str(out_path)
            fig_idx += 1

    return ref_to_path


# ── Chunking ───────────────────────────────────────────────────────────────────

def _build_chunks(
    dl_doc: DoclingDocument,
    source_name: str,
    ref_to_path: dict[str, str],
    text_splitter: RecursiveCharacterTextSplitter,
) -> list[Document]:
    """
    Convert a parsed DoclingDocument into LangChain Document chunks.

    Uses Docling's HybridChunker with a tiktoken tokenizer to produce
    semantically meaningful chunks that respect document structure (headings,
    paragraphs, tables). Falls back to the plain RecursiveCharacterTextSplitter
    on the raw markdown export if the HybridChunker fails.

    Two chunk types are produced:
    - Text chunks (chunk_type='text'): split further by text_splitter if too long
    - Figure chunks (chunk_type='figure'): kept whole, with image_path in metadata

    Figures that the HybridChunker produced no surrounding text for get a
    stub placeholder so they still exist in the vector store.
    """
    # Use tiktoken for accurate token counting during chunking
    fast_tok = OpenAITokenizer(
        tokenizer=tiktoken.encoding_for_model("gpt-3.5-turbo"),
        max_tokens=8192,
    )
    chunker = HybridChunker(tokenizer=fast_tok)

    # Try HybridChunker first; fall back to simple splitter on failure
    try:
        raw_chunks = list(chunker.chunk(dl_doc))
    except Exception as e:
        print(f"        HybridChunker failed ({e}), falling back to RecursiveCharacterTextSplitter")
        raw_text = dl_doc.export_to_markdown()
        fallback = text_splitter.create_documents(
            [raw_text],
            metadatas=[{"source": source_name, "chunk_type": "text"}],
        )
        # Add figure stubs so images aren't lost on fallback
        for img_path in ref_to_path.values():
            fallback.append(Document(
                page_content=f"[Abbildung aus {source_name}]",
                metadata={"source": source_name, "chunk_type": "figure", "image_path": img_path},
            ))
        return fallback

    text_chunks:   list[Document] = []
    figure_chunks: list[Document] = []

    for raw_chunk in raw_chunks:
        text = chunker.contextualize(raw_chunk)
        if not text.strip():
            continue

        # Check whether any item in this chunk is a visual element
        image_path: str | None = None
        for doc_item in raw_chunk.meta.doc_items:
            if doc_item.label in _VISUAL_LABELS:
                image_path = ref_to_path.get(doc_item.self_ref)
                if image_path:
                    break

        if image_path:
            figure_chunks.append(Document(
                page_content=text,
                metadata={"source": source_name, "chunk_type": "figure", "image_path": image_path},
            ))
        else:
            text_chunks.append(Document(
                page_content=text,
                metadata={"source": source_name, "chunk_type": "text"},
            ))

    # Find figure refs that the chunker already produced context text for
    chunked_refs = {
        doc_item.self_ref
        for raw_chunk in raw_chunks
        for doc_item in raw_chunk.meta.doc_items
        if doc_item.label in _VISUAL_LABELS
    }

    # Add stub chunks for figures with no surrounding text
    for ref, img_path in ref_to_path.items():
        if ref not in chunked_refs:
            figure_chunks.append(Document(
                page_content=f"[Abbildung aus {source_name}]",
                metadata={"source": source_name, "chunk_type": "figure", "image_path": img_path},
            ))

    # Split long text chunks into smaller pieces
    split_text = text_splitter.split_documents(text_chunks)
    for c in split_text:
        c.metadata.setdefault("chunk_type", "text")

    return split_text + figure_chunks


def _extract_docx_links(file_path: str, source_name: str) -> list[Document]:
    """
    Extract all hyperlinks from a DOCX file and return them as a single text chunk.
    Links embedded in Word documents are otherwise invisible to Docling's chunker,
    so this ensures they remain searchable in the vector store.
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    links: list[tuple[str, str]] = []
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            for hyperlink in paragraph._element.iter(qn("w:hyperlink")):
                r_id = hyperlink.get(qn("r:id"))
                if r_id and r_id in doc.part.rels:
                    rel = doc.part.rels[r_id]
                    if "hyperlink" in rel.reltype:
                        url  = rel.target_ref
                        text = "".join(t.text for t in hyperlink.iter(qn("w:t"))).strip()
                        if url:
                            links.append((text, url))
    except Exception:
        pass

    if not links:
        return []

    link_lines = "\n".join(
        f"- {text}: {url}" if text else f"- {url}" for text, url in links
    )
    return [Document(
        page_content=f"Verweise und Links in {source_name}:\n{link_lines}",
        metadata={"source": source_name, "chunk_type": "text"},
    )]


# ── Worker process ─────────────────────────────────────────────────────────────
# Docling parsing runs in a process pool for parallelism. Each worker gets its
# own converter instance (initialised once per worker via the initializer function)
# to avoid pickling the converter across process boundaries.

_WORKER_CONVERTER: DocumentConverter | None = None


def _worker_init() -> None:
    """Initialise the Docling converter in each worker process (OCR disabled)."""
    global _WORKER_CONVERTER
    _WORKER_CONVERTER = _make_converter()


def _worker_init_ocr() -> None:
    """Initialise the Docling converter in each worker process (OCR enabled).
    Used for the second-pass retry of scanned PDFs that produced 0 chunks."""
    global _WORKER_CONVERTER
    pdf_opts        = PdfPipelineOptions()
    pdf_opts.do_ocr = True
    _WORKER_CONVERTER = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts)}
    )


def _process_one(
    path_str: str,
    image_store_str: str,
    chunk_size: int,
    chunk_overlap: int,
) -> tuple[str, list[dict], str | None]:
    """
    Parse a single file inside a worker process.

    Returns a tuple of:
    - filename (str)
    - list of serialised chunk dicts (page_content + metadata) — empty on failure
    - error string if something went wrong, None on success

    Chunks are serialised to plain dicts because LangChain Document objects
    can't always be pickled cleanly across process boundaries.
    """
    path = Path(path_str)
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n## ", "\n### ", "\n\n", "\n", " ", ""],
        )
        assert _WORKER_CONVERTER is not None, "worker not initialized"

        result = _WORKER_CONVERTER.convert(str(path))
        dl_doc = result.document

        # Include the file extension in the stem so that figures from files with
        # the same name but different extensions (e.g. report.xlsx vs report.pptx)
        # don't overwrite each other in the image store
        safe_stem = f"{path.stem}_{path.suffix.lstrip('.')}"
        ref_map   = _extract_images(dl_doc, safe_stem, Path(image_store_str))
        chunks    = _build_chunks(dl_doc, path.name, ref_map, splitter)

        if path.suffix.lower() == ".docx":
            chunks.extend(_extract_docx_links(str(path), path.name))

        return (
            path.name,
            [{"page_content": c.page_content, "metadata": c.metadata} for c in chunks],
            None,
        )
    except Exception as e:
        return (path.name, [], f"{type(e).__name__}: {e}")


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    # ── Step 0: convert legacy file formats ───────────────────────────────────
    print("[0/3] Converting legacy office files...")
    _convert_legacy(DOCS_DIR)

    # ── Step 0b: convert WMF/EMF images inside Office ZIPs ───────────────────
    print("[0b/3] Converting WMF/EMF images inside cached Office files...")
    wmf_total = _batch_convert_vectors_in_cache()
    if wmf_total:
        print(f"      Converted {wmf_total} WMF/EMF image(s) to PNG.")
    else:
        print("      No WMF/EMF images to convert.")

    # Collect all files in the converted cache (skip hidden files and Office lock files)
    paths = [
        p for p in CONVERTED_CACHE.rglob("*")
        if p.is_file()
        and not p.name.startswith(".")
        and not p.name.startswith("~$")
        and not p.name.endswith("~")
    ]
    if not paths:
        print(f"No files found in {CONVERTED_CACHE}")
        return

    # ── Step 1: parallel Docling parsing (OCR off) ────────────────────────────
    print(f"[1/3] Parsing {len(paths)} file(s) with Docling — {NUM_WORKERS} workers, OCR off...")

    all_chunks:       list[Document]    = []
    zero_chunk_paths: dict[str, bool]   = {}  # files that parsed OK but produced 0 chunks (likely scanned PDFs)
    failed_files:     list[str]         = []

    from pebble import ProcessPool
    from concurrent.futures import TimeoutError

    with ProcessPool(max_workers=NUM_WORKERS, initializer=_worker_init) as pool:
        # Schedule all files; each has a 3-minute timeout per file
        futures = {
            pool.schedule(
                _process_one,
                args=(str(p), config.IMAGE_STORE_PATH, config.CHUNK_SIZE, config.CHUNK_OVERLAP),
                timeout=180,
            ): p
            for p in paths
        }
        for i, future in enumerate(as_completed(futures), 1):
            p = futures[future]
            try:
                name, chunk_dicts, error = future.result()
            except TimeoutError:
                name, chunk_dicts, error = p.name, [], "TimeoutError: killed after 3 minutes"
            except Exception as e:
                name, chunk_dicts, error = p.name, [], f"Worker crash: {type(e).__name__} - {e}"

            if error:
                print(f"      [{i}/{len(paths)}] FAILED {name}: {error}", flush=True)
                failed_files.append(name)
            else:
                if len(chunk_dicts) == 0:
                    zero_chunk_paths[name] = True  # flag for OCR retry
                for d in chunk_dicts:
                    all_chunks.append(Document(page_content=d["page_content"], metadata=d["metadata"]))
                print(f"      [{i}/{len(paths)}] {name} ({len(chunk_dicts)} chunks)", flush=True)

    # Move permanently failed files to a review folder and log them
    if failed_files:
        failed_dir = Path("./data/failed_files")
        failed_dir.mkdir(parents=True, exist_ok=True)
        for fname in failed_files:
            for p in paths:
                if p.name == fname:
                    try:
                        dest = failed_dir / p.name
                        p.replace(dest)
                        print(f"      Moved failed file {p.name} to {dest}")
                    except Exception as e:
                        print(f"      Could not move {p.name}: {e}")
        with open(failed_dir / "failed_files.txt", "w") as f:
            for fname in failed_files:
                f.write(fname + "\n")
        print(f"      All failed files moved to {failed_dir}/ and listed in failed_files.txt")

    text_count   = sum(1 for c in all_chunks if c.metadata["chunk_type"] == "text")
    figure_count = sum(1 for c in all_chunks if c.metadata["chunk_type"] == "figure")
    print(f"      Total: {text_count} text chunks, {figure_count} figure chunks")
    print(f"      Images saved to {config.IMAGE_STORE_PATH}/")

    # ── Step 1b: retry scanned PDFs with OCR enabled ──────────────────────────
    # Any PDF that produced 0 chunks on the first pass is likely a scanned document
    # with no machine-readable text layer. Re-run it with OCR to extract text.
    ocr_paths = [
        p for p in paths
        if p.suffix.lower() == ".pdf" and zero_chunk_paths.get(p.name)
    ]
    if ocr_paths:
        print(f"[1b/3] Re-processing {len(ocr_paths)} scanned PDF(s) with OCR enabled...")
        with ProcessPool(max_workers=max(1, NUM_WORKERS // 2), initializer=_worker_init_ocr) as pool:
            futures2 = {
                pool.schedule(
                    _process_one,
                    args=(str(p), config.IMAGE_STORE_PATH, config.CHUNK_SIZE, config.CHUNK_OVERLAP),
                    timeout=300,  # OCR is slower — allow 5 minutes per file
                ): p
                for p in ocr_paths
            }
            for i, future in enumerate(as_completed(futures2), 1):
                p = futures2[future]
                try:
                    name, chunk_dicts, error = future.result()
                except TimeoutError:
                    name, chunk_dicts, error = p.name, [], "TimeoutError: killed after 5 minutes"
                except Exception as e:
                    name, chunk_dicts, error = p.name, [], f"Worker crash: {type(e).__name__} - {e}"

                if error:
                    print(f"      [{i}/{len(ocr_paths)}] FAILED {name}: {error}", flush=True)
                else:
                    for d in chunk_dicts:
                        all_chunks.append(Document(page_content=d["page_content"], metadata=d["metadata"]))
                    print(f"      [{i}/{len(ocr_paths)}] {name} ({len(chunk_dicts)} chunks via OCR)", flush=True)

    # ── Step 2: load embedding model ──────────────────────────────────────────
    # BAAI/bge-m3 produces dense embeddings for semantic search.
    # Batch size is larger on GPU to make use of parallelism.
    print("[2/3] Loading embedding model...")
    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-m3",
        model_kwargs={"device": config.EMBED_DEVICE},
        encode_kwargs={"batch_size": 64 if config.EMBED_DEVICE == "cuda" else 16},
    )

    # ── Step 3: embed and index everything into Qdrant ────────────────────────
    # force_recreate=True wipes the existing collection so re-runs start clean.
    # Hybrid mode combines dense (BGE-M3) + sparse (BM25) vectors for better recall.
    print("[3/3] Embedding and indexing into Qdrant...")
    QdrantVectorStore.from_documents(
        all_chunks,
        embedding=embeddings,
        sparse_embedding=FastEmbedSparse(model_name="Qdrant/bm25"),
        retrieval_mode=RetrievalMode.HYBRID,
        path=config.QDRANT_PATH,
        collection_name=config.COLLECTION,
        force_recreate=True,
    )
    print(f"      Done — {len(all_chunks)} total chunks in '{config.COLLECTION}'")


if __name__ == "__main__":
    main()