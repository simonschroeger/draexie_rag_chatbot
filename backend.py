"""
DRÄXIE API — FastAPI + LangChain RAG pipeline.
Run with: uvicorn backend:app --reload
"""

import base64
import io
import json
import math
import os
import re
import sqlite3
import tempfile
from pathlib import Path

import httpx
from fastapi import BackgroundTasks, FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from docling.datamodel.base_models import DocItemLabel
from docling.datamodel.document import DoclingDocument, PictureItem
from docling.document_converter import DocumentConverter
from docling_core.transforms.chunker import HybridChunker
from langchain_community.chat_message_histories import SQLChatMessageHistory
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_qdrant import FastEmbedSparse, QdrantVectorStore, RetrievalMode
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from PIL import Image as PILImage

import config


# ── Constants ─────────────────────────────────────────────────────────────────

# Labels that identify visual elements (figures/charts) inside parsed documents
_VISUAL_LABELS = {DocItemLabel.PICTURE, DocItemLabel.CHART}

# Image filtering thresholds applied at ingestion time.
# Images failing either check are skipped entirely — never saved or described.
_IMG_MIN_DIM    = 150   # skip if BOTH width AND height are below this (icons, bullet symbols)
_IMG_MAX_ASPECT = 10.0  # skip if longest side / shortest side exceeds this (decorative lines)

# Relevance gate for the off-topic guard.
# If the best matching text chunk scores below this threshold, the LLM is skipped
# entirely and the "nothing found" answer is returned immediately.
# Tune down to allow more questions through; tune up to be stricter.
_MIN_TEXT_RELEVANCE = 0.40

# How many figure chunks to attach to a response at most, and the minimum
# relevance score a figure chunk must reach to be included at all.
_MAX_FIGURE_RESULTS = 3
_MIN_FIGURE_SCORE   = 0.45


# ── Databases ─────────────────────────────────────────────────────────────────

# SQLite connection strings / paths
CONV_DB = "sqlite:///conversations.db"   # stores full conversation history (LangChain managed)
META_DB = "draxie.db"                    # stores feedback ratings and ingestion logs


def init_db() -> None:
    """Create the metadata tables on first startup if they don't exist yet."""
    con = sqlite3.connect(META_DB)
    con.executescript("""
        CREATE TABLE IF NOT EXISTS feedback (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT,
            rating          TEXT,
            question        TEXT,
            answer          TEXT,
            created_at      DATETIME DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS ingestion_log (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            filename    TEXT,
            chunks      INTEGER,
            status      TEXT,
            error       TEXT,
            created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
        );
    """)
    con.commit()
    con.close()


# Run once at import time so tables always exist before any request arrives
init_db()


def get_session_history(session_id: str) -> SQLChatMessageHistory:
    """Return the LangChain message history object for a given conversation ID."""
    return SQLChatMessageHistory(session_id=session_id, connection=CONV_DB)


def _ensure_conv_tables(con: sqlite3.Connection) -> None:
    """
    Create the two extra conversation tables if they don't exist yet.
    These live in conversations.db (managed by LangChain) rather than draxie.db,
    so they need their own creation guard.
    """
    con.execute("""
        CREATE TABLE IF NOT EXISTS message_sources (
            conversation_id TEXT NOT NULL,
            msg_index       INTEGER NOT NULL,
            sources         TEXT,   -- JSON list of source filenames
            chunks          TEXT,   -- JSON list of chunk objects shown to the user
            PRIMARY KEY (conversation_id, msg_index)
        )
    """)
    con.execute("""
        CREATE TABLE IF NOT EXISTS conversation_titles (
            conversation_id TEXT PRIMARY KEY,
            title           TEXT NOT NULL,
            created_at      DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)


def _save_message_sources(conv_id: str, msg_idx: int, sources: list, chunks: list) -> None:
    """
    Persist which source documents and chunk excerpts were used for a specific
    assistant message. This lets the frontend re-display sources when loading
    an old conversation.
    """
    con = sqlite3.connect("conversations.db")
    _ensure_conv_tables(con)
    con.execute(
        "INSERT OR REPLACE INTO message_sources (conversation_id, msg_index, sources, chunks) VALUES (?,?,?,?)",
        (conv_id, msg_idx, json.dumps(sources), json.dumps(chunks)),
    )
    con.commit()
    con.close()


# ── Document ingestion helpers ─────────────────────────────────────────────────
# These functions are shared between create_database.py (bulk indexing)
# and the live POST /documents upload endpoint.

# One Docling converter instance, reused across all ingestion calls
_converter = DocumentConverter()

# Text splitter used to break large text chunks into smaller pieces
_splitter = RecursiveCharacterTextSplitter(
    chunk_size=config.CHUNK_SIZE,
    chunk_overlap=config.CHUNK_OVERLAP,
    separators=["\n## ", "\n### ", "\n\n", "\n", " ", ""],
)


def _extract_images(
    dl_doc: DoclingDocument,
    source_stem: str,
) -> dict[str, str]:
    """
    Walk every item in a parsed document and save qualifying PictureItems as PNG files.

    Filtering rules (applied before saving):
    - Skip images where both width AND height are below _IMG_MIN_DIM (icons, bullets)
    - Skip images with an extreme aspect ratio above _IMG_MAX_ASPECT (decorative lines)

    Returns a dict mapping each item's self_ref ID to its saved file path.
    """
    image_store = Path(config.IMAGE_STORE_PATH)
    image_store.mkdir(parents=True, exist_ok=True)
    ref_to_path: dict[str, str] = {}
    fig_idx = 0

    for item, _level in dl_doc.iterate_items():
        if isinstance(item, PictureItem):
            try:
                img = item.get_image(dl_doc)
            except Exception:
                fig_idx += 1
                continue

            if img is None:
                fig_idx += 1
                continue

            w, h = img.size

            if w == 0 or h == 0:
                fig_idx += 1
                continue

            # Skip tiny decorative elements (icons, bullet points, small logos)
            if w < _IMG_MIN_DIM and h < _IMG_MIN_DIM:
                fig_idx += 1
                continue

            # Skip extreme-aspect-ratio images (horizontal/vertical decorative lines)
            if max(w, h) / min(w, h) > _IMG_MAX_ASPECT:
                fig_idx += 1
                continue

            out_path = image_store / f"{source_stem}__fig{fig_idx:04d}.png"
            img.save(out_path, "PNG")
            ref_to_path[item.self_ref] = str(out_path)
            fig_idx += 1

    return ref_to_path


def _build_chunks(
    dl_doc: DoclingDocument,
    source_name: str,
    ref_to_path: dict[str, str],
) -> list[Document]:
    """
    Convert a parsed DoclingDocument into a list of LangChain Document chunks.

    Two types of chunks are produced:
    - Text chunks (chunk_type='text'): passed through the text splitter
    - Figure chunks (chunk_type='figure'): kept whole, with image_path in metadata.
      Each figure chunk gets a German vision description generated by the LLM so
      it becomes searchable by text queries.

    Any figure that Docling's HybridChunker produced no surrounding text for
    gets a stub placeholder so it still appears in the vector store.
    """
    chunker = HybridChunker()
    text_docs: list[Document]   = []
    figure_docs: list[Document] = []

    for raw_chunk in chunker.chunk(dl_doc):
        text = chunker.contextualize(raw_chunk)
        if not text.strip():
            continue

        # Check if any item in this chunk is a figure/chart
        image_path: str | None = None
        for doc_item in raw_chunk.meta.doc_items:
            if doc_item.label in _VISUAL_LABELS:
                image_path = ref_to_path.get(doc_item.self_ref)
                if image_path:
                    break

        if image_path:
            figure_docs.append(Document(
                page_content=text,
                metadata={"source": source_name, "chunk_type": "figure", "image_path": image_path},
            ))
        else:
            text_docs.append(Document(
                page_content=text,
                metadata={"source": source_name, "chunk_type": "text"},
            ))

    # Find which figure refs the chunker already covered
    chunked_refs = {
        doc_item.self_ref
        for raw_chunk in chunker.chunk(dl_doc)
        for doc_item in raw_chunk.meta.doc_items
        if doc_item.label in _VISUAL_LABELS
    }

    # Add stub chunks for figures the chunker produced no text context for
    for ref, img_path in ref_to_path.items():
        if ref not in chunked_refs:
            figure_docs.append(Document(
                page_content=f"[Abbildung aus {source_name}]",
                metadata={"source": source_name, "chunk_type": "figure", "image_path": img_path},
            ))

    # Generate a German vision description for each figure so it's retrievable by text search
    for doc in figure_docs:
        img_path = doc.metadata.get("image_path", "")
        if not img_path:
            continue
        surrounding = doc.page_content if not doc.page_content.startswith("[Abbildung") else ""
        description = _describe_image(img_path, surrounding)
        if description:
            doc.page_content = description
            doc.metadata["image_description"] = description

    # Split long text chunks into smaller pieces
    split_text = _splitter.split_documents(text_docs)
    for c in split_text:
        c.metadata.setdefault("chunk_type", "text")

    return split_text + figure_docs


def _extract_docx_links(file_path: str, source_name: str) -> list[Document]:
    """
    Extract all hyperlinks from a DOCX file and return them as a single text chunk.
    This ensures links buried in Word documents remain retrievable via the RAG search.
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    links: list[tuple[str, str]] = []
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            for hyperlink in paragraph._element.iter(qn("w:hyperlink")):
                r_id = hyperlink.get(qn("r:id"))
                if r_id and r_id in doc.part.rels:
                    rel = doc.part.rels[r_id]
                    if "hyperlink" in rel.reltype:
                        url  = rel.target_ref
                        text = "".join(t.text for t in hyperlink.iter(qn("w:t"))).strip()
                        if url:
                            links.append((text, url))
    except Exception:
        pass

    if not links:
        return []

    link_lines = "\n".join(
        f"- {text}: {url}" if text else f"- {url}" for text, url in links
    )
    return [Document(
        page_content=f"Verweise und Links in {source_name}:\n{link_lines}",
        metadata={"source": source_name, "chunk_type": "text"},
    )]


# ── Vision helpers ─────────────────────────────────────────────────────────────

def _image_to_b64(image_path: str) -> str | None:
    """
    Load an image from disk, resize it to fit within the configured token budget,
    and return it as a base64-encoded PNG string ready to send to the vision model.
    Returns None if the image can't be loaded.
    """
    try:
        img = PILImage.open(image_path).convert("RGB")
        max_pixels = config.VISION_TOKEN_BUDGET * 512
        w, h = img.size
        if w * h > max_pixels:
            scale = math.sqrt(max_pixels / (w * h))
            img = img.resize((max(1, int(w * scale)), max(1, int(h * scale))), PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode()
    except Exception as e:
        print(f"[image_to_b64] FAILED {image_path}: {e}", flush=True)
        return None


def _describe_image(image_path: str, surrounding_text: str = "") -> str:
    """
    Send an image to the LLM (vision mode) and get back a German description
    suitable for indexing. The surrounding document text is included as context
    so the model can produce a more relevant description.
    Returns an empty string on failure.
    """
    b64 = _image_to_b64(image_path)
    if not b64:
        return ""

    ctx = (
        f"\n\nUmliegender Textkontext aus dem Dokument:\n{surrounding_text[:800]}"
        if surrounding_text.strip() else ""
    )
    prompt = (
        "Beschreibe dieses Bild auf Deutsch ausführlich für ein technisches Dokumentensystem. "
        "Nenne: Art der Visualisierung, alle sichtbaren Zahlen und Bezeichnungen, "
        "den wesentlichen Inhalt und welche Schlussfolgerung es vermittelt."
        + ctx
    )

    try:
        response = llm.invoke([HumanMessage(content=[
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            {"type": "text",      "text": prompt},
        ])])
        return response.content.strip()
    except Exception as e:
        print(f"[describe_image] FAILED {image_path}: {e}", flush=True)
        return ""


# ── RAG pipeline setup ────────────────────────────────────────────────────────
# These objects are created once at startup and reused for every request.

# Embedding model: converts text to dense vectors for semantic search
_embeddings = HuggingFaceEmbeddings(
    model_name="BAAI/bge-m3",
    model_kwargs={"device": config.EMBED_DEVICE},
)

# Qdrant vector store client (local file-based, no separate server needed)
_qdrant_client = QdrantClient(path=config.QDRANT_PATH)

# Vector store: combines dense embeddings + sparse BM25 for hybrid search
vectorstore = QdrantVectorStore(
    client=_qdrant_client,
    collection_name=config.COLLECTION,
    embedding=_embeddings,
    sparse_embedding=FastEmbedSparse(model_name="Qdrant/bm25"),
    retrieval_mode=RetrievalMode.HYBRID,
)

# The local LLM served via Ollama
llm = ChatOllama(
    base_url=config.OLLAMA_BASE_URL,
    model=config.LLM_MODEL,
    temperature=0.3,
    num_ctx=config.LLM_NUM_CTX,
)

# Basic retriever: fetches the top-K most relevant chunks for a query
base_retriever = vectorstore.as_retriever(search_kwargs={"k": config.TOP_K})

# Query expansion chain: the LLM rewrites a user question into 3 alternative
# phrasings (2 in the same language, 1 translated) to improve retrieval recall
_expand_chain = (
    ChatPromptTemplate.from_template(
        "You are a multilingual query expansion assistant for a bilingual (German/English) document corpus.\n"
        "Detect the language of the search query below.\n"
        "Generate exactly 3 alternative phrasings following this rule:\n"
        "  - If the query is in German: produce 2 alternatives in German and 1 alternative translated into English.\n"
        "  - If the query is in English: produce 2 alternatives in English and 1 alternative translated into German.\n"
        "The cross-language alternative must be the last element.\n"
        "Return ONLY a valid JSON array of exactly 3 strings, with no extra text:\n"
        "[\"alternative1\", \"alternative2\", \"cross_language_alternative\"]\n\n"
        "Query: {question}"
    )
    | llm
    | StrOutputParser()
)


def _parse_expand_alternatives(raw: str) -> list[str]:
    """
    Parse the JSON array of alternative queries from the LLM's raw output.
    Falls back to treating each non-empty line as an alternative if JSON parsing fails.
    """
    start = raw.find("[")
    end   = raw.rfind("]")
    if start == -1 or end == -1:
        return []
    try:
        items = json.loads(raw[start:end + 1])
        return [s.strip() for s in items if isinstance(s, str) and s.strip()]
    except Exception:
        return [ln.strip() for ln in raw.splitlines() if ln.strip()]


# Follow-up suggestion chain: after answering, the LLM generates 3 questions
# the user might want to ask next, based on the retrieved context and the answer
_suggest_chain = (
    ChatPromptTemplate.from_template(
        "You are a DRÄXLMAIER sales onboarding assistant.\n"
        "Based on the context and answer below, generate exactly 3 short follow-up questions "
        "in the same language as the answer, that the user might ask next and that are answerable from the context.\n"
        "Return ONLY a JSON array, e.g.: [\"Question 1?\", \"Question 2?\", \"Question 3?\"]\n\n"
        "Context:\n{context}\n\nAnswer:\n{answer}"
    )
    | ChatOllama(
        base_url=config.OLLAMA_BASE_URL,
        model=config.LLM_MODEL,
        temperature=0.1,
    )
    | StrOutputParser()
)


def _parse_suggestions(raw: str) -> list[str]:
    """Parse the JSON array of follow-up suggestions from the LLM's raw output."""
    start = raw.find("[")
    end   = raw.rfind("]")
    if start == -1 or end == -1:
        return []
    try:
        items = json.loads(raw[start:end + 1])
        return [s for s in items if isinstance(s, str)][:3]
    except Exception:
        return []


def multi_retrieve(question: str) -> tuple[list, list]:
    """
    Retrieve relevant document chunks for a user question using multi-query expansion.

    Steps:
    1. The LLM expands the question into 3 alternative phrasings
    2. All 4 queries (original + 3 alternatives) run against the vector store
    3. Text chunks are deduplicated and returned as the ground-truth sources
    4. Figure chunks are retrieved separately with a relevance score gate —
       only images that genuinely match the query above _MIN_FIGURE_SCORE are included

    Returns (text_docs, figure_docs).
    """
    raw_expand   = _expand_chain.invoke({"question": question})
    alternatives = _parse_expand_alternatives(raw_expand)
    queries      = [question] + alternatives   # 4 queries total

    # Collect unique text chunks across all query variants
    seen, text_docs = set(), []
    for q in queries:
        for doc in base_retriever.invoke(q):
            if doc.metadata.get("chunk_type") == "figure":
                continue  # figures are retrieved separately below
            if doc.page_content not in seen:
                seen.add(doc.page_content)
                text_docs.append(doc)

    # Retrieve figure chunks with a score threshold — silently skip if Qdrant filter fails
    figure_docs: list = []
    try:
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        fig_filter = Filter(must=[
            FieldCondition(key="metadata.chunk_type", match=MatchValue(value="figure"))
        ])
        fig_results = vectorstore.similarity_search_with_relevance_scores(
            query=question,
            k=20,
            filter=fig_filter,
        )
        for doc, score in fig_results:
            if score <= _MIN_FIGURE_SCORE:
                continue
            img = doc.metadata.get("image_path", "")
            if not img or not Path(img).exists():
                continue
            figure_docs.append(doc)
            if len(figure_docs) >= _MAX_FIGURE_RESULTS:
                break
    except Exception:
        pass

    return text_docs, figure_docs


def format_docs(docs: list) -> str:
    """Format text chunks into a numbered context string to pass to the LLM."""
    return "\n\n".join(f"[{i+1}] {d.page_content}" for i, d in enumerate(docs))


# ── Response formatting ────────────────────────────────────────────────────────
# Gemma sometimes produces malformed output when combining bullet lists with citations.
# These patterns fix the most common issues before streaming to the frontend.

_PIPE_BEFORE_CITE  = re.compile(r'\s*\|\s*(?=\[\d)')        # "text | [1]" → "text [1]"
_JAMMED_CITES      = re.compile(r'\](\[\d+\])')              # "][2]" → "] [2]"
_LONE_CITE_PERIOD  = re.compile(r'((?:\[\d+\]\s*)+)\.\s*$', re.MULTILINE)  # "[1]." → "[1]"


def repair_gemma_formatting(text: str) -> str:
    """Fix Gemma's broken bullet+pipe citation patterns before streaming to the frontend."""
    lines = text.split('\n')
    result: list[str] = []
    in_fence = False
    for line in lines:
        if line.strip().startswith('```'):
            in_fence = not in_fence
            result.append(line)
            continue
        if not in_fence:
            line = _PIPE_BEFORE_CITE.sub(' ', line)
        result.append(line)
    text = '\n'.join(result)
    text = _JAMMED_CITES.sub(r'] \1', text)
    text = _LONE_CITE_PERIOD.sub(lambda m: m.group(1).rstrip(), text)
    return text


# ── System prompt ─────────────────────────────────────────────────────────────
# Defines DRÄXIE's persona, grounding rules, citation format, and language behaviour.
# This is sent as the system message on every chat request.

_SYSTEM_PROMPT = """\
You are DRÄXIE, a precise onboarding assistant for DRÄXLMAIER sales staff.

GROUNDING RULES — ABSOLUTE, NO EXCEPTIONS
- You have NO general knowledge. Your training data does not exist in this context. The ONLY facts you may state are those explicitly present in the document excerpts provided below.
- NEVER answer from memory, training data, or world knowledge — even if you are certain of the answer. This applies to every topic: people, companies, history, science, politics, geography, or anything else.
- If the question cannot be answered from the excerpts alone, you MUST use the no-answer sentence. No partial guesses. No "based on general knowledge". Silence on everything outside the excerpts.
- Every sentence that states a fact MUST end with its source marker [N] (e.g. [1], [2]). No exceptions. If a sentence draws on multiple excerpts, cite each separately with no space: [1][3]. Never write [1, 3] or [1,3].
- Use only sequential [1], [2], [3] … markers — they correspond to the numbered Quellen entries shown to the user.
- When asked for specific numbers, deadlines, names, or lists, state them directly and exactly as they appear in the documents.
- Some excerpts are marked [Abbildung aus …]. These indicate a figure or diagram exists in the source. You may cite them with [N] to point the user to the visual, but do not invent or describe their contents — only text excerpts are authoritative.

LINKS AND REFERENCES
- If the document excerpts contain hyperlinks or references to other documents, always include them verbatim in your answer so the user can follow them directly.

WHEN THE DOCUMENTS ARE INCOMPLETE OR OFF-TOPIC
- If the question is about a person, event, organization, or topic that does not appear anywhere in the excerpts, respond IMMEDIATELY with only the no-answer sentence. Do not attempt to answer.
- If the documents fully answer the question, answer completely.
- If the documents only partially answer it, give what the documents support, then clearly state what is missing. If relevant links to further resources appear in the excerpts, include them.
- If the documents do not answer the question at all, use this exact sentence (German): "Dazu habe ich in den verfügbaren Unterlagen leider nichts gefunden." — do not paraphrase it.
  English equivalent: "I couldn't find anything on that in the available documents."
- Do not include source citations in a no-answer response.

FORBIDDEN WORDS — never use these in any response, in any language:
- "Kontext" → use "die Unterlagen", "die Dokumente", or "die Auszüge" instead
- "Prompt" → never say this to users
- "Token" → never say this to users
- "Embedding" → never say this to users
- "Chunk" / "Chunks" → use "Abschnitt" or "Auszug" if referring to a passage
- "Query" (as a technical term) → use "Ihre Frage" or "Ihre Anfrage"
- "Retrieval" → never say this to users
- "Vektor" / "Vector" (in the AI/search sense) → never say this to users
Speak like a knowledgeable colleague, not a software system.

PLAIN LANGUAGE
Use domain-specific terminology (Fahrzeugprogramm, Baureihe, OEM, Sonderausstattung, etc.) when appropriate
for DRÄXLMAIER sales staff. Never use IT or computing terminology (algorithms, APIs, databases, etc.) —
the FORBIDDEN WORDS list above covers the most common cases.
Write like a knowledgeable automotive colleague, not a software system.

FORMAT
- Default to clear prose. Use a table or bullet list only when the content is genuinely a list, comparison, or set of steps — not for single facts.
- NEVER use raw markdown pipe-table syntax (lines containing "|"). It will not render correctly.
- When the answer is best shown as a table, emit a fenced block with this EXACT syntax (nothing else inside the fence):
  ```ui-table
    {{"title": "optional title", "columns": ["Col A", "Col B"], "rows": [["Val 1", "Val 2"], ["Val 3", "Val 4"]]}}
  ```
  Put normal prose before and after the block. Only use it for genuine tabular data.
- For glossary-style answers with multiple terms and definitions, use plain bullet list format: **Term:** definition. Do NOT use a table for this.
- [N] markers correspond to the numbered document excerpts provided.
- NEVER place a "|" pipe character inside a bullet point or numbered list item. A line starting with "-" or "1." must contain only prose text — never table columns.

CITATION PLACEMENT
- [N] markers go ONLY at the very end of a complete sentence. Correct: "Bauteile werden verwendet. [1]" — Wrong: "Bauteile [1] werden verwendet." or "| [1]. text |"
- NEVER place [N] at the start or middle of a sentence, or at the start of any content block.
- When multiple citations follow each other, write them with no period between: "[1][2]" not "[1]. [2]" or "[1][2].".
- A lone citation with only a period ("| [10].") is forbidden — every citation must follow a complete sentence of prose.

MULTI-DEFINITION FORMAT
When a term has multiple distinct meanings, ALWAYS use this numbered list format — never a table, never mixed bullet+pipe:
1. **Bedeutung A:** Vollständiger Erklärungssatz in Prosa. [N]
2. **Bedeutung B:** Vollständiger Erklärungssatz in Prosa. [N]

ANTWORTFORMAT — BEISPIELE (STRIKT EINHALTEN)

FALSCH — so NICHT (Aufzählung mit Pipe-Spalten und Zitationen mittendrin):
- **COP:** Erklärung Teil 1 | [1][3][7]. Erklärung Teil 2 | [1][8]. Erklärung Teil 3 | [8].
- **Conformity of Production:** [2][3].

RICHTIG — so JA (nummerierte Liste, vollständige Sätze, Zitationen am Satzende):
Der Begriff **COP** hat folgende Bedeutungen:

1. **Carry Over Part (Übernahmeteil):** Bauteile, die unverändert in verschiedenen Produkten verwendet werden, aber keine Normteile sind. [1] Sie senken Entwicklungskosten und sind Kern des Plattformkonzepts. [3]

2. **Conformity of Production:** Übereinstimmung der Serienproduktion mit dem genehmigten Typ. [2]

3. **Central Ordering Process (KOVP):** Zentraler Bestellprozess bei DRÄXLMAIER. [3]

REGEL: Jedes Listenelement enthält vollständige Sätze. Zitationen [N] stehen IMMER am Ende eines vollständigen Satzes — niemals am Satzanfang, niemals zwischen Wörtern, niemals nach einem Pipe-Zeichen.

SINGLE ANSWER
Give one single direct answer. Do not list multiple possibilities when you can give a clear answer.
Express uncertainty once and clearly — never format it as a bullet list of options.

LANGUAGE
- Always reply in the same language as the user's question. German question → German answer. English question → English answer.

CONVERSATION
- Use earlier messages in the conversation for context when relevant.
"""

# The answer prompt: system instructions + conversation history + current question with context
_answer_prompt = ChatPromptTemplate.from_messages([
    ("system", _SYSTEM_PROMPT),
    MessagesPlaceholder(variable_name="history"),
    ("human", "Unterlagen:\n{context}\n\nANWEISUNG: Beantworte ausschließlich auf Basis der Unterlagen oben. Wenn das Thema der Frage dort nicht vorkommt, antworte NUR mit: \"Dazu habe ich in den verfügbaren Unterlagen leider nichts gefunden.\" Kein Allgemeinwissen, keine eigene Recherche.\n\nFrage: {question}"),
])

# LangChain chain that automatically loads and saves conversation history from SQLite
_chain_with_history = RunnableWithMessageHistory(
    _answer_prompt | llm | StrOutputParser(),
    get_session_history,
    input_messages_key="question",
    history_messages_key="history",
)


# ── No-answer detection ────────────────────────────────────────────────────────
# The LLM sometimes paraphrases "I don't know" in different ways.
# These helpers normalise any such response to a single canonical sentence.

_NO_ANSWER_DE = "Dazu habe ich in den verfügbaren Unterlagen leider nichts gefunden."
_NO_ANSWER_EN = "I couldn't find anything on that in the available documents."

# Lowercase trigger phrases that indicate the LLM couldn't find an answer
_NO_ANSWER_TRIGGERS = [
    "keine informationen", "nicht gefunden", "nicht in den unterlagen",
    "nicht in den bereitgestellten", "nicht in meinen", "leider nichts",
    "keine angaben", "nicht enthalten", "nicht verfügbar", "nicht vorhanden",
    "i couldn't find", "no information", "not found in", "cannot find",
    "not available in", "not mentioned in", "no relevant", "nothing found",
    "keine relevanten", "nicht erwähnt", "nicht behandelt",
]

# Common English function words used to detect the question language
_EN_MARKERS = {"what", "who", "where", "when", "why", "how", "is", "are",
               "was", "were", "does", "did", "can", "could", "would", "tell",
               "explain", "describe", "give", "show", "list", "find", "the", "a"}


def _detect_language(text: str) -> str:
    """
    Heuristically detect whether a text is English or German.
    Counts word overlap with known English and German marker sets.
    Returns 'en' or 'de'.
    """
    words   = set(re.findall(r"[a-zA-ZäöüÄÖÜß]+", text.lower()))
    en_hits = len(words & _EN_MARKERS)
    de_hits = len(words & {"wie", "welche", "welcher", "welches", "warum",
                            "wann", "können", "gibt", "sind", "haben", "bitte",
                            "erkläre", "zeige", "nenne", "beschreibe"})
    return "en" if en_hits > de_hits else "de"


def _normalize_no_answer(text: str, question: str) -> str:
    """
    Replace any LLM paraphrase of "I don't know" with the canonical no-answer sentence.
    If the response already contains citation markers like [1] or [2], it found real
    content and is left untouched.
    """
    if re.search(r'\[\d+\]', text):
        return text  # real answer with citations — don't touch
    lower = text.lower()
    if any(t in lower for t in _NO_ANSWER_TRIGGERS):
        return _NO_ANSWER_EN if _detect_language(question) == "en" else _NO_ANSWER_DE
    return text


def _best_text_score(question: str) -> float:
    """
    Check the relevance score of the single best-matching text chunk for a question.
    Used as a fast off-topic gate before calling the LLM — if nothing in the vector
    store is close enough to the question, we skip the LLM entirely.
    Fails open (returns 1.0) if Qdrant can't be reached, so the LLM still runs.
    """
    try:
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        results = vectorstore.similarity_search_with_relevance_scores(
            question, k=1,
            filter=Filter(must=[FieldCondition(key="metadata.chunk_type", match=MatchValue(value="text"))])
        )
        return results[0][1] if results else 0.0
    except Exception:
        return 1.0


# ── FastAPI app ───────────────────────────────────────────────────────────────

app = FastAPI(title="DRÄXIE API")

# Allow all origins so the frontend can talk to the backend during local development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve the built React frontend's static assets
app.mount("/assets", StaticFiles(directory="frontend/dist/assets"), name="assets")

# Serve extracted document images (referenced by image_url in chunk metadata)
_images_dir = Path(config.IMAGE_STORE_PATH)
_images_dir.mkdir(parents=True, exist_ok=True)
app.mount("/images", StaticFiles(directory=str(_images_dir)), name="images")

# Serve original uploaded files so the in-app document viewer can open them
_uploads_dir = Path("./data/uploads")
_uploads_dir.mkdir(parents=True, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=str(_uploads_dir)), name="uploads")


# ── Request/response models ───────────────────────────────────────────────────

class ChatRequest(BaseModel):
    question:        str
    conversation_id: str


class FeedbackRequest(BaseModel):
    conversation_id: str
    rating:          str   # "up" or "down"
    question:        str
    answer:          str


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    """Serve the React SPA entry point."""
    return FileResponse("frontend/dist/index.html")


@app.get("/draexie-mascot.js")
async def mascot_js():
    """Serve the animated mascot JavaScript bundle."""
    return FileResponse("frontend/dist/draexie-mascot.js", media_type="application/javascript")


@app.post("/chat")
def chat(req: ChatRequest):
    """
    Main chat endpoint. Returns a Server-Sent Events stream with three event types:
    1. sources/chunks  — which documents and excerpts were retrieved
    2. token           — the full generated answer (sent as one piece after generation)
    3. done            — signals completion, includes follow-up suggestions

    Two code paths depending on whether relevant figures were found:
    - Multimodal path: images are base64-encoded and sent to the LLM alongside the text context
    - Text-only path:  standard LangChain chain with conversation history
    """
    # Off-topic gate: skip the LLM entirely if nothing relevant is in the vector store
    best_score = _best_text_score(req.question)
    if best_score < _MIN_TEXT_RELEVANCE:
        no_ans = _NO_ANSWER_EN if _detect_language(req.question) == "en" else _NO_ANSWER_DE
        def _nope():
            yield f"data: {json.dumps({'sources': [], 'chunks': []})}\n\n"
            yield f"data: {json.dumps({'token': no_ans})}\n\n"
            yield f"data: {json.dumps({'done': True, 'suggestions': []})}\n\n"
        return StreamingResponse(_nope(), media_type="text/event-stream")

    # Retrieve relevant text and figure chunks
    text_docs, figure_docs = multi_retrieve(req.question)

    # Source filenames (deduplicated, text docs only — these are shown to the user)
    sources = list(dict.fromkeys(
        Path(d.metadata.get("source", "unknown")).name for d in text_docs
    ))

    # Chunk objects sent to the frontend for the "sources" panel
    all_docs = text_docs + figure_docs
    chunks = [
        {
            "num":       i + 1,
            "source":    Path(d.metadata.get("source", "unknown")).name,
            "text":      d.page_content[:500] if d.metadata.get("chunk_type") != "figure" else "",
            "image_url": (
                f"/images/{Path(d.metadata['image_path']).name}"
                if d.metadata.get("chunk_type") == "figure"
                   and d.metadata.get("image_path")
                   and Path(d.metadata["image_path"]).exists()
                else None
            ),
        }
        for i, d in enumerate(all_docs)
    ]

    # Text context sent to the LLM (figures are never included here — only their descriptions were indexed)
    context = format_docs(text_docs)

    # Collect unique image file paths for the multimodal path
    image_paths: list[str] = []
    seen_paths:  set[str]  = set()
    for d in figure_docs:
        p = d.metadata.get("image_path", "")
        if p and p not in seen_paths and Path(p).exists():
            image_paths.append(p)
            seen_paths.add(p)

    def generate():
        # Always send sources first so the frontend can display them immediately
        yield f"data: {json.dumps({'sources': sources, 'chunks': chunks})}\n\n"

        if image_paths:
            # ── Multimodal path: question + images sent directly to the LLM ──
            b64_images = [b for p in image_paths if (b := _image_to_b64(p)) is not None]

            # Load conversation history, capped at last 10 messages to stay within context window
            history_obj  = get_session_history(req.conversation_id)
            history_msgs = history_obj.messages[-10:]

            # Build the multimodal message: images first, then text (optimal for Gemma)
            human_content: list[dict] = [
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b}"}}
                for b in b64_images
            ] + [
                {"type": "text", "text": (
                    f"Kontext:\n{context}\n\n"
                    "ANWEISUNG: Beantworte ausschließlich auf Basis der Unterlagen oben. "
                    "Wenn das Thema der Frage dort nicht vorkommt, antworte NUR mit: "
                    "\"Dazu habe ich in den verfügbaren Unterlagen leider nichts gefunden.\" "
                    f"Kein Allgemeinwissen, keine eigene Recherche.\n\nFrage: {req.question}"
                )},
            ]

            messages = [
                SystemMessage(content=_SYSTEM_PROMPT),
                *history_msgs,
                HumanMessage(content=human_content),
            ]

            # Stream the response, collect the full answer, then send it as one token event
            full_answer = ""
            for chunk in llm.stream(messages):
                tok = chunk.content if hasattr(chunk, "content") else str(chunk)
                if tok:
                    full_answer += tok

            full_answer = _normalize_no_answer(repair_gemma_formatting(full_answer), req.question)
            yield f"data: {json.dumps({'token': full_answer})}\n\n"

            # Manually persist the turn since we bypassed LangChain's history chain
            history_obj.add_user_message(req.question)
            history_obj.add_ai_message(full_answer)
            _save_message_sources(req.conversation_id, len(history_obj.messages) - 1, sources, chunks)

        else:
            # ── Text-only path: standard LangChain chain with automatic history ──
            full_answer = ""
            for token in _chain_with_history.stream(
                {"context": context, "question": req.question},
                config={"configurable": {"session_id": req.conversation_id}},
            ):
                full_answer += token

            full_answer = _normalize_no_answer(repair_gemma_formatting(full_answer), req.question)
            yield f"data: {json.dumps({'token': full_answer})}\n\n"

            # Save which sources were used for this message
            history_after = get_session_history(req.conversation_id)
            _save_message_sources(req.conversation_id, len(history_after.messages) - 1, sources, chunks)

        # Generate follow-up question suggestions based on the answer and context
        suggestions: list[str] = []
        try:
            raw         = _suggest_chain.invoke({"context": context, "answer": full_answer})
            suggestions = _parse_suggestions(raw)
        except Exception:
            pass

        yield f"data: {json.dumps({'done': True, 'suggestions': suggestions})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ── Conversation history routes ───────────────────────────────────────────────

@app.get("/conversations")
def list_conversations():
    """Return all conversations, sorted newest first, with title and message count."""
    db_path = Path("conversations.db")
    if not db_path.exists():
        return []

    con = sqlite3.connect(db_path)
    _ensure_conv_tables(con)
    rows = con.execute("""
        SELECT m.session_id, MIN(m.id) as first_id, MAX(m.id) as last_id,
               COUNT(*) as cnt, t.title
        FROM message_store m
        LEFT JOIN conversation_titles t ON t.conversation_id = m.session_id
        GROUP BY m.session_id
        ORDER BY last_id DESC
    """).fetchall()
    con.close()

    result = []
    for session_id, first_id, last_id, cnt, stored_title in rows:
        if stored_title:
            title = stored_title
        else:
            # Fall back to the first user message as the title
            history = get_session_history(session_id)
            msgs = history.messages
            if not msgs:
                continue
            title = next(
                (m.content[:80] for m in msgs if isinstance(m, HumanMessage)),
                "Unbenanntes Gespräch",
            )
        result.append({
            "id":            session_id,
            "title":         title,
            "created_at":    str(last_id),
            "message_count": cnt,
        })
    return result


@app.post("/conversations/{conversation_id}/title")
def generate_title(conversation_id: str):
    """
    Ask the LLM to generate a short German title for a conversation based on
    its first user message, then persist it in conversation_titles.
    """
    history     = get_session_history(conversation_id)
    first_human = next(
        (m.content for m in history.messages if isinstance(m, HumanMessage)),
        None,
    )
    if not first_human:
        return {"title": "Unbenanntes Gespräch"}

    try:
        resp  = llm.invoke([HumanMessage(content=(
            f"Erstelle einen kurzen deutschen Titel (maximal 6 Wörter, keine Anführungszeichen) "
            f"für ein Gespräch, das mit dieser Frage beginnt:\n\n{first_human[:300]}"
        ))])
        title = resp.content.strip().strip('"').strip("'")[:80] or first_human[:60]
    except Exception:
        title = first_human[:60]

    con = sqlite3.connect("conversations.db")
    _ensure_conv_tables(con)
    con.execute(
        "INSERT OR REPLACE INTO conversation_titles (conversation_id, title) VALUES (?,?)",
        (conversation_id, title),
    )
    con.commit()
    con.close()
    return {"title": title}


@app.get("/conversations/search")
def search_conversations(q: str):
    """
    Full-text search across conversations. Searches three places:
    1. Message content
    2. Conversation titles
    3. Source document filenames cited in answers
    Returns one result per matching conversation with a short excerpt.
    """
    db_path = Path("conversations.db")
    if not db_path.exists() or not q.strip():
        return []

    con    = sqlite3.connect(db_path)
    _ensure_conv_tables(con)
    like_q = f"%{q}%"
    seen:   dict[str, dict] = {}

    def _msg_text(raw: str) -> str:
        """Extract plain text from LangChain's JSON message storage format."""
        try:
            return json.loads(raw).get("data", {}).get("content", raw)
        except Exception:
            return raw

    def _make_entry(
        session_id: str, text: str, stored_title: str | None,
        msg_idx: int, created_at: str | None = None,
    ) -> dict | None:
        """Build a search result entry with a highlighted excerpt around the match."""
        idx = text.lower().find(q.lower())
        if idx == -1:
            return None
        start   = max(0, idx - 55)
        end     = min(len(text), idx + len(q) + 55)
        excerpt = ("…" if start > 0 else "") + text[start:end] + ("…" if end < len(text) else "")
        return {
            "id":          session_id,
            "title":       stored_title or text[:80],
            "excerpt":     excerpt,
            "match_start": idx - start + (1 if start > 0 else 0),
            "match_len":   len(q),
            "msg_index":   msg_idx,
            "created_at":  created_at,
        }

    # 1. Search message content
    for session_id, raw_msg, stored_title, msg_idx, created_at in con.execute("""
        SELECT m.session_id, m.message, t.title,
               (SELECT COUNT(*) FROM message_store m2
                WHERE m2.session_id = m.session_id AND m2.id <= m.id) - 1,
               t.created_at
        FROM message_store m
        LEFT JOIN conversation_titles t ON t.conversation_id = m.session_id
        WHERE m.message LIKE ?
        ORDER BY m.id DESC
    """, (like_q,)).fetchall():
        if session_id not in seen:
            entry = _make_entry(session_id, _msg_text(raw_msg), stored_title, int(msg_idx or 0), created_at)
            if entry:
                seen[session_id] = entry

    # 2. Search conversation titles
    for session_id, title, created_at in con.execute(
        "SELECT conversation_id, title, created_at FROM conversation_titles WHERE title LIKE ?", (like_q,)
    ).fetchall():
        if session_id not in seen:
            entry = _make_entry(session_id, title, title, 0, created_at)
            if entry:
                seen[session_id] = entry

    # 3. Search source document names used in answers
    for session_id, sources_json, msg_idx, stored_title, created_at in con.execute("""
        SELECT ms.conversation_id, ms.sources, ms.msg_index, t.title, t.created_at
        FROM message_sources ms
        LEFT JOIN conversation_titles t ON t.conversation_id = ms.conversation_id
        WHERE ms.sources LIKE ?
        ORDER BY ms.msg_index DESC
    """, (like_q,)).fetchall():
        if session_id not in seen:
            try:
                readable = ", ".join(json.loads(sources_json))
            except Exception:
                readable = sources_json
            if not stored_title:
                row = con.execute(
                    "SELECT message FROM message_store WHERE session_id=? ORDER BY id LIMIT 1",
                    (session_id,),
                ).fetchone()
                stored_title = (_msg_text(row[0])[:80] if row else None)
            entry = _make_entry(session_id, readable, stored_title, int(msg_idx or 0), created_at)
            if entry:
                seen[session_id] = entry

    con.close()
    return list(seen.values())


@app.get("/conversations/{conversation_id}")
def get_conversation(conversation_id: str):
    """
    Return the full message history for a conversation, with sources and chunk
    metadata attached to each assistant message.
    """
    history = get_session_history(conversation_id)

    # Load which sources were cited for each assistant message
    sources_map: dict[int, dict] = {}
    db_path = Path("conversations.db")
    if db_path.exists():
        con = sqlite3.connect(db_path)
        try:
            rows = con.execute(
                "SELECT msg_index, sources, chunks FROM message_sources WHERE conversation_id=?",
                (conversation_id,),
            ).fetchall()
            for idx, src_json, chk_json in rows:
                sources_map[int(idx)] = {
                    "sources": json.loads(src_json) if src_json else [],
                    "chunks":  json.loads(chk_json) if chk_json else [],
                }
        except sqlite3.OperationalError:
            pass  # table doesn't exist yet in older databases
        finally:
            con.close()

    messages = []
    for i, m in enumerate(history.messages):
        msg: dict = {
            "role":    "user" if isinstance(m, HumanMessage) else "assistant",
            "content": m.content if isinstance(m.content, str) else "",
        }
        if i in sources_map:
            msg["sources"] = sources_map[i]["sources"]
            msg["chunks"]  = sources_map[i]["chunks"]
        messages.append(msg)

    return {"id": conversation_id, "messages": messages}


@app.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: str):
    """Delete all messages for a conversation from the history store."""
    get_session_history(conversation_id).clear()
    return {"deleted": conversation_id}


# ── Feedback ──────────────────────────────────────────────────────────────────

@app.post("/feedback")
def feedback(req: FeedbackRequest):
    """Store a thumbs-up or thumbs-down rating for an answer in the metadata DB."""
    con = sqlite3.connect(META_DB)
    con.execute(
        "INSERT INTO feedback (conversation_id, rating, question, answer) VALUES (?,?,?,?)",
        (req.conversation_id, req.rating, req.question, req.answer),
    )
    con.commit()
    con.close()
    return {"ok": True}


# ── Document upload ───────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".pptx", ".xlsx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp"}


def _ingest_file(tmp_path: str, original_name: str) -> None:
    """
    Background task: parse an uploaded file, embed it, and add it to the vector store.
    Also saves a copy of the original file so the document viewer can serve it.
    Logs the result (success or error) to the ingestion_log table.
    """
    # Keep original file accessible for the in-app document viewer
    upload_copy = _uploads_dir / original_name
    upload_copy.write_bytes(Path(tmp_path).read_bytes())

    con = sqlite3.connect(META_DB)
    try:
        result  = _converter.convert(tmp_path)
        dl_doc  = result.document
        ref_map = _extract_images(dl_doc, Path(original_name).stem)
        chunks  = _build_chunks(dl_doc, original_name, ref_map)
        if Path(original_name).suffix.lower() == ".docx":
            chunks.extend(_extract_docx_links(str(upload_copy), original_name))
        vectorstore.add_documents(chunks)
        con.execute(
            "INSERT INTO ingestion_log (filename, chunks, status) VALUES (?,?,?)",
            (original_name, len(chunks), "ok"),
        )
    except Exception as exc:
        con.execute(
            "INSERT INTO ingestion_log (filename, chunks, status, error) VALUES (?,?,?,?)",
            (original_name, 0, "error", str(exc)),
        )
    finally:
        con.commit()
        con.close()
        os.unlink(tmp_path)   # clean up the temp file regardless of success/failure


@app.post("/documents")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    """
    Accept a document upload, save it to a temp file, and kick off background ingestion.
    Returns immediately — the client can poll /documents/status to track progress.
    """
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return {"error": f"Unsupported file type: {suffix}"}

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(await file.read())
    tmp.close()

    background_tasks.add_task(_ingest_file, tmp.name, file.filename)
    return {"status": "processing", "filename": file.filename}


@app.get("/documents/status")
def ingestion_status():
    """Return the 20 most recent ingestion log entries (newest first)."""
    con  = sqlite3.connect(META_DB)
    rows = con.execute(
        "SELECT filename, chunks, status, error, created_at FROM ingestion_log ORDER BY id DESC LIMIT 20"
    ).fetchall()
    con.close()
    return [
        {"filename": r[0], "chunks": r[1], "status": r[2], "error": r[3], "created_at": r[4]}
        for r in rows
    ]


# ── Health check ──────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    """
    Returns the operational status of all three system components:
    - Qdrant: checks the collection exists and reports chunk count
    - Ollama: pings the local model server
    - Last ingestion: reports the most recently successfully indexed file
    """
    # Check Qdrant
    try:
        count    = _qdrant_client.count(config.COLLECTION).count
        qdrant_ok = True
    except Exception:
        count    = 0
        qdrant_ok = False

    # Check Ollama
    try:
        r        = httpx.get(f"{config.OLLAMA_BASE_URL}/api/tags", timeout=3.0)
        ollama_ok = r.status_code == 200
    except Exception:
        ollama_ok = False

    # Fetch last successful ingestion from the log
    con  = sqlite3.connect(META_DB)
    last = con.execute(
        "SELECT filename, created_at FROM ingestion_log WHERE status='ok' ORDER BY id DESC LIMIT 1"
    ).fetchone()
    con.close()

    return {
        "status": "ok" if (qdrant_ok and ollama_ok) else "degraded",
        "qdrant": {"ok": qdrant_ok, "chunks": count, "collection": config.COLLECTION},
        "ollama": {"ok": ollama_ok, "model": config.LLM_MODEL},
        "last_ingestion": {"filename": last[0], "at": last[1]} if last else None,
    }


# ── SPA fallback ──────────────────────────────────────────────────────────────

@app.get("/{full_path:path}", include_in_schema=False)
async def spa_fallback(full_path: str):
    """
    Catch-all route: serve the React SPA for any URL that isn't an API endpoint.
    This allows client-side routing (e.g. /chat/123) to work when the page is refreshed.
    """
    return FileResponse("frontend/dist/index.html")